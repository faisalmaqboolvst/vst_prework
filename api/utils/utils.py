import docx
from PIL import Image
import pytesseract
from io import BytesIO


# Helper functions to extract text and images from .docx
def extract_text(docx_path):
    """Extracts plain text from a .docx file."""
    doc = docx.Document(docx_path)
    full_text = [para.text for para in doc.paragraphs]
    return "\n".join(full_text)


def extract_images(docx_path):
    """Extracts images from a .docx file and returns PIL image objects."""
    doc = docx.Document(docx_path)
    images = []
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_data = rel.target_part.blob  # Image binary data
            image = Image.open(BytesIO(image_data))
            images.append(image)
    return images


def perform_ocr(images):
    """Performs OCR on a list of images and returns the extracted text."""
    ocr_texts = [pytesseract.image_to_string(image) for image in images]
    return ocr_texts


# Helper function to serialize MongoDB ObjectId
def serialize_data(data):
    """Converts MongoDB documents to a JSON-serializable format."""
    if isinstance(data, list):
        for item in data:
            if "_id" in item:
                item["_id"] = str(item["_id"])
    elif "_id" in data:
        data["_id"] = str(data["_id"])
    return data
